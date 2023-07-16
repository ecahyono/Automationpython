import docx

# Buat objek dokumen Word baru
doc = docx.Document()

# Tambahkan judul instrumen Litmas
judul = doc.add_paragraph("INSTRUMEN LITMAS")
judul.bold = True
judul.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
judul.runs[0].underline = True
judul.runs[0].font.size = docx.shared.Pt(14)

# Tambahkan konten instrumen Litmas
petunjuk = doc.add_paragraph("Petunjuk: ...")
tugas = doc.add_paragraph("Tugas: ...")
evaluasi = doc.add_paragraph("Evaluasi: ...")

# Simpan dokumen dengan path dan nama file yang diinginkan
doc.save('C:\\Users\\user\\Documents\\TRCH\\Automationpython\\SDP\\SDP_TAHAP2\\Bapas\\Pendaftaran\\Register\\nama_file.docx')
